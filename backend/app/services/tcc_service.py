import io
import json
import logging
import re

from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.llm_client import generate_text

load_dotenv()

logger = logging.getLogger(__name__)

STRUCTURE_SYSTEM_PROMPT = """Você é um assistente acadêmico especializado em analisar a estrutura de Trabalhos de Conclusão de Curso (TCC).

Analise o texto abaixo e identifique as seções presentes. Retorne APENAS um JSON válido (sem markdown):

{
  "secoes": [
    {
      "titulo": "título da seção",
      "tipo": "capa|resumo|introducao|referencial_teorico|metodologia|resultados|discussao|conclusao|referencias|anexo|outro",
      "pagina_estimada": null,
      "completude": "completa|parcial|ausente",
      "sugestoes": ["sugestão 1", "sugestão 2"]
    }
  ],
  "estrutura_geral": "descrição geral da estrutura encontrada",
  "secoes_ausentes": ["seção que deveria existir mas não foi encontrada"],
  "recomendacoes_estrutura": ["recomendação 1", "recomendação 2"]
}

Regras:
- Liste APENAS seções que você identificar com confiança no texto.
- Se uma seção existe mas está incompleta, marque completude como "parcial".
- Sugestões devem ser específicas e acionáveis, nunca reescrever o texto do aluno.
- Seja rigoroso: se não houver evidência de uma seção, não a invente."""

REVIEW_SYSTEM_PROMPT = """Você é um revisor acadêmico especializado em clareza e coesão textual para TCC.

Analise o texto abaixo e identifique problemas de clareza e coesão. 
ATENÇÃO: NUNCA reescreva o texto do aluno. Apenas aponte problemas.

Retorne APENAS um JSON válido (sem markdown):

{
  "problemas": [
    {
      "trecho": "trecho exato do texto com problema (máx 200 chars)",
      "tipo": "clareza|coesao|coerencia|concordancia|repeticao|ambiguidade|pontuacao|registro",
      "gravidade": "alta|media|baixa",
      "sugestao_generica": "descrição do problema sem reescrever o texto. Ex: 'Esta frase pode ser simplificada' ao invés de 'Escreva: X'"
    }
  ],
  "resumo_geral": "resumo (2-3 frases) dos principais problemas encontrados",
  "pontos_fortes": ["aspecto positivo 1", "aspecto positivo 2"]
}

Regras:
- NUNCA reescreva o texto do aluno.
- Se não houver problemas significativos, retorne {"problemas": [], "resumo_geral": "Texto claro e coeso.", "pontos_fortes": [...]}
- Limite a no máximo 10 problemas."""

REFERENCES_ABNT_SYSTEM_PROMPT = """Você é um especialista em normas ABNT para trabalhos acadêmicos.

Analise o texto abaixo e extraia as referências bibliográficas, verificando sua conformidade com a ABNT NBR 6023.

Retorne APENAS um JSON válido (sem markdown):

{
  "referencias": [
    {
      "texto_extraido": "referência completa como aparece no texto",
      "elementos_obrigatorios": {
        "autor": "presente|ausente|parcial",
        "titulo": "presente|ausente|parcial",
        "edicao": "presente|ausente|parcial|nao_se_aplica",
        "local": "presente|ausente",
        "editora": "presente|ausente",
        "ano": "presente|ausente"
      },
      "conforme_abnt": true|false,
      "problemas": ["problema 1", "problema 2"],
      "sugestao_correcao": "descrição de como ajustar, sem reescrever"
    }
  ],
  "total_referencias": 0,
  "conformidade_geral": "percentual aproximado de conformidade (0-100)",
  "recomendacoes": ["recomendação geral 1", "recomendação geral 2"]
}

Regras:
- Extraia TODAS as referências que encontrar.
- Se um elemento é obrigatório mas não se aplica (ex: edição para artigo), use "nao_se_aplica".
- Seja criterioso: uma referência sem ano está incompleta mesmo que tudo mais esteja correto."""


def _try_parse_json(content: str) -> dict:
    content = content.strip()
    if content.startswith("```"):
        content = re.sub(r"^```\w*\n?", "", content)
        content = re.sub(r"\n?```$", "", content)
    try:
        return json.loads(content)
    except json.JSONDecodeError as e:
        logger.warning(f"JSON parse failed at {e.pos}: {e.msg}, attempting recovery")
        truncated = content[:e.pos]
        if not truncated.rstrip().endswith('"') and '"' in truncated:
            truncated = truncated.rsplit('"', 1)[0] + '"\n'
        open_braces = truncated.count("{") - truncated.count("}")
        open_brackets = truncated.count("[") - truncated.count("]")
        truncated += "}" * open_braces
        truncated += "]" * open_brackets
        try:
            return json.loads(truncated)
        except json.JSONDecodeError:
            raise ValueError(f"Não foi possível recuperar o JSON da resposta: {content[:200]}...")


async def analyze_structure(texto: str) -> dict:
    if len(texto) > 500000:
        logger.warning(f"Texto truncado de {len(texto)} para 500000 caracteres para análise de estrutura")

    content = await generate_text(
        system_prompt=STRUCTURE_SYSTEM_PROMPT,
        user_text=f"Texto do TCC:\n\n{texto[:500000]}",
        max_tokens=8192,
        temperature=0.3,
    )
    return _try_parse_json(content)


async def review_text(texto: str) -> dict:
    if len(texto) > 500000:
        logger.warning(f"Texto truncado de {len(texto)} para 500000 caracteres para revisão textual")

    content = await generate_text(
        system_prompt=REVIEW_SYSTEM_PROMPT,
        user_text=f"Texto do TCC:\n\n{texto[:500000]}",
        max_tokens=8192,
        temperature=0.3,
    )
    return _try_parse_json(content)


async def extract_references(texto: str) -> dict:
    if len(texto) > 500000:
        logger.warning(f"Texto truncado de {len(texto)} para 500000 caracteres para extração de referências")

    content = await generate_text(
        system_prompt=REFERENCES_ABNT_SYSTEM_PROMPT,
        user_text=f"Texto do TCC:\n\n{texto[:500000]}",
        max_tokens=8192,
        temperature=0.3,
    )
    return _try_parse_json(content)


def create_combined_report(
    estrutura: dict,
    revisao: dict,
    referencias: dict,
    plano: str = "free",
) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    def _heading(text: str, level: int = 1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.size = Pt(14 if level == 1 else 12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

    def _para(text: str):
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(6)

    def _bullet(text: str):
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.first_line_indent = Cm(1.25)

    title = doc.add_heading("RELATÓRIO DE ANÁLISE — TCC", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

    _heading("1. ESTRUTURA DO TRABALHO")
    secs = estrutura.get("secoes", [])
    for sec in secs:
        _para(f"{sec.get('titulo', '?')} ({sec.get('tipo', '?')}) — {sec.get('completude', '?')}")
        suggs = sec.get("sugestoes", [])
        for s in suggs:
            _bullet(s)
    _para(f"Estrutura geral: {estrutura.get('estrutura_geral', '')}")
    ausentes = estrutura.get("secoes_ausentes", [])
    if ausentes:
        _para("Seções ausentes:")
        for s in ausentes:
            _bullet(s)
    for rec in estrutura.get("recomendacoes_estrutura", []):
        _bullet(rec)

    _heading("2. REVISÃO TEXTUAL")
    problemas = revisao.get("problemas", [])
    if not problemas:
        _para("Nenhum problema significativo encontrado.")
    else:
        gravidade_order = {"alta": 0, "media": 1, "baixa": 2}
        problemas.sort(key=lambda p: gravidade_order.get(p.get("gravidade", "baixa"), 3))
        for prob in problemas:
            trecho = prob.get("trecho", "")[:150]
            tipo = prob.get("tipo", "")
            gravidade = prob.get("gravidade", "")
            _para(f"[{gravidade.upper()}] [{tipo}] \"{trecho}\"")
            sugestao = prob.get("sugestao_generica", "")
            if sugestao:
                _bullet(sugestao)
    pontos = revisao.get("pontos_fortes", [])
    if pontos:
        _para("Pontos fortes:")
        for p in pontos:
            _bullet(p)

    _heading("3. REFERÊNCIAS ABNT")
    refs = referencias.get("referencias", [])
    _para(f"Total de referências encontradas: {referencias.get('total_referencias', len(refs))}")
    conf = referencias.get("conformidade_geral", 0)
    _para(f"Conformidade geral ABNT: {conf}%")
    for ref in refs:
        texto = ref.get("texto_extraido", "")
        conforme = ref.get("conforme_abnt", False)
        status = "✓" if conforme else "✗"
        _para(f"{status} {texto[:200]}")
        for prob in ref.get("problemas", []):
            _bullet(prob)
    for rec in referencias.get("recomendacoes", []):
        _bullet(rec)

    footer_text = "Gerado com Trilha — Inteligência Artificial para seus estudos"
    if plano == "free":
        footer_text += " | Plano Gratuito"
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(footer_text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
