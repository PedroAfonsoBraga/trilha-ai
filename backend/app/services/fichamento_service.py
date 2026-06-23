import io
import json
import logging
import os
import re

from dotenv import load_dotenv
import httpx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()

logger = logging.getLogger(__name__)

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
DEEPSEEK_URL = "https://api.deepseek.com/v1/chat/completions"
DEFAULT_MODEL = "deepseek-v4-flash"

FICHAMENTO_ABNT_SYSTEM = """Você é um especialista em fichamento acadêmico no formato ABNT.

Analise o texto abaixo e produza um fichamento estruturado. Retorne um JSON com:

{
  "referencia": "SOBRENOME, Nome. Título (se identificável). Órgão/Instituição, data.",
  "tema": "tema central do texto",
  "objetivo": "objetivo principal",
  "metodologia": "metodologia ou abordagem (se identificável)",
  "principais_pontos": ["ponto 1", "ponto 2", "ponto 3", ...],
  "citacoes_relevantes": ["citação 1", "citação 2"],
  "conclusao": "conclusão ou considerações finais",
  "comentarios": "análise crítica breve (2-3 frases)"
}

Regras:
- A referencia deve seguir ABNT: SOBRENOME, Nome. Título. Cidade: Editora, ano.
- Se não for possível identificar algum campo, use "Não identificado".
- Principais pontos deve ter de 5 a 10 itens.
- Citações relevantes deve ser transcrições literais do texto.

Responda APENAS com o JSON, sem markdown ou texto adicional."""


def _try_parse_json(content: str) -> dict:
    content = content.strip()
    if content.startswith("```"):
        content = re.sub(r"^```\w*\n?", "", content)
        content = re.sub(r"\n?```$", "", content)
    try:
        return json.loads(content)
    except json.JSONDecodeError as e:
        # Try to salvage: close unclosed strings/braces
        logger.warning(f"JSON parse failed at {e.pos}: {e.msg}, attempting recovery")
        truncated = content[:e.pos]
        # Close any unclosed strings
        if truncated.rstrip().endswith('"') is False and '"' in truncated:
            truncated = truncated.rsplit('"', 1)[0] + '"\n'
        # Close unclosed arrays/objects
        open_braces = truncated.count("{") - truncated.count("}")
        open_brackets = truncated.count("[") - truncated.count("]")
        truncated += "}" * open_braces
        truncated += "]" * open_brackets
        try:
            return json.loads(truncated)
        except json.JSONDecodeError:
            raise ValueError(f"Não foi possível recuperar o JSON da resposta: {content[:200]}...")


async def gerar_fichamento_ia(texto: str) -> dict:
    if not DEEPSEEK_API_KEY:
        raise ValueError("DEEPSEEK_API_KEY não configurada")

    # Truncate input to leave room for response (80000 tokens ≈ 320KB)
    input_text = texto[:120000]

    async with httpx.AsyncClient(timeout=180.0) as client:
        response = await client.post(
            DEEPSEEK_URL,
            headers={
                "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": DEFAULT_MODEL,
                "messages": [
                    {"role": "system", "content": FICHAMENTO_ABNT_SYSTEM},
                    {"role": "user", "content": f"Texto para fichamento:\n\n{input_text}"},
                ],
                "max_tokens": 8192,
                "temperature": 0.3,
            },
        )
        response.raise_for_status()
        data = response.json()

    content = data["choices"][0]["message"]["content"]
    finish_reason = data["choices"][0].get("finish_reason", "")

    if finish_reason == "length":
        logger.warning("DeepSeek response truncated by token limit")

    return _try_parse_json(content)


def criar_docx_fichamento(fichamento: dict, plano: str = "free") -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    title = doc.add_heading("FICHAMENTO", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

    sections = [
        ("REFERÊNCIA", fichamento.get("referencia", "Não identificado")),
        ("TEMA", fichamento.get("tema", "Não identificado")),
        ("OBJETIVO", fichamento.get("objetivo", "Não identificado")),
        ("METODOLOGIA", fichamento.get("metodologia", "Não identificado")),
    ]

    for label, value in sections:
        h = doc.add_heading(label, level=2)
        for run in h.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        p = doc.add_paragraph(value)
        p.paragraph_format.first_line_indent = Cm(1.25)

    doc.add_heading("PRINCIPAIS PONTOS", level=2)
    for ponto in fichamento.get("principais_pontos", []):
        p = doc.add_paragraph(ponto, style="List Bullet")
        p.paragraph_format.first_line_indent = Cm(1.25)

    doc.add_heading("CITAÇÕES RELEVANTES", level=2)
    for citacao in fichamento.get("citacoes_relevantes", []):
        p = doc.add_paragraph(f'"{citacao}"')
        p.paragraph_format.left_indent = Cm(4)
        p.runs[0].font.italic = True

    doc.add_heading("CONCLUSÃO", level=2)
    p = doc.add_paragraph(fichamento.get("conclusao", "Não identificado"))
    p.paragraph_format.first_line_indent = Cm(1.25)

    doc.add_heading("COMENTÁRIOS", level=2)
    p = doc.add_paragraph(fichamento.get("comentarios", "Não identificado"))
    p.paragraph_format.first_line_indent = Cm(1.25)

    # Footer branding
    footer_text = "Gerado com Trilha — Inteligência Artificial para seus estudos"
    if plano == "free":
        footer_text += " | Plano Gratuito"

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(footer_text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
